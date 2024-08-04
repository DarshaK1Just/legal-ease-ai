from pdfreader import PDFDocument
import docx

def read_pdf(file):
    text = ""
    try:
        with open(file, 'rb') as f:
            pdf = PDFDocument(f)
            for page in pdf.pages:
                text += page.text()
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def read_docx(file):
    doc = docx.Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def read_txt(file):
    return file.read().decode("utf-8")

def read_file(file):
    if file.type == "application/pdf":
        return read_pdf(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx(file)
    elif file.type == "text/plain":
        return read_txt(file)
    else:
        raise ValueError("Unsupported file type")
